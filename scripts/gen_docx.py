# -*- coding: utf-8 -*-
"""Markdown + Mermaid PNG → Word 文档生成器（mermaid-doc-renderer skill）

把 markdown 文档转换成 .docx，满足：
  - 无 markdown 语法残留（剥离 **加粗**、`代码`、链接、表格语法）
  - mermaid 代码块用已渲染的 PNG 图嵌入
  - 标题/表格/列表/引用/代码块保留语义格式

用法：
  python gen_docx.py <md_path> --docx <out.docx> --png-dir <渲染PNG目录> [--title "标题"] [--subtitle "副标题"]

依赖：
  - Python：python-docx, Pillow
"""
import os, re, sys, argparse
from PIL import Image

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C_BLUE = RGBColor(0x1F, 0x4E, 0x79)
C_GRAY = RGBColor(0x60, 0x60, 0x60)
C_CODE = RGBColor(0x00, 0x40, 0x20)
MAX_W = 6.3   # 英寸
MAX_H = 6.3


def strip_md_inline(text):
    """把行内 markdown 拆成 (plain_text, is_bold, is_mono) 片段列表"""
    tokens = []
    pattern = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append((text[pos:m.start()], False, False))
        tok = m.group(0)
        if tok.startswith('**'):
            inner = tok[2:-2]
            inner = re.sub(r'`([^`]*)`', r'\1', inner)
            tokens.append((inner, True, False))
        else:
            tokens.append((tok[1:-1], False, True))
        pos = m.end()
    if pos < len(text):
        tokens.append((text[pos:], False, False))
    return tokens


def build_doc(md_path, docx_path, png_dir, title='', subtitle=''):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def set_east(run, east):
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), east)

    def add_run(p, text, east='宋体', size=10.5, bold=None, color=None, mono=False):
        run = p.add_run(text)
        run.font.size = Pt(size)
        if mono:
            run.font.name = 'Consolas'
            set_east(run, '等线')
        else:
            run.font.name = 'Calibri'
            set_east(run, east)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = color
        return run

    def add_para(tokens, size=10.5, align=None, space_after=6, east='宋体'):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if isinstance(tokens, str):
            tokens = [(tokens, False, False)]
        for t, b, mono in tokens:
            if mono:
                add_run(p, t, east='等线', size=size, bold=False, mono=True)
            else:
                add_run(p, t, east=east, size=size, bold=b or None)
        return p

    def add_heading(text, level):
        h = doc.add_heading(level=level)
        for r in list(h.runs):
            r.text = ''
        sizes = {0: 22, 1: 17, 2: 14, 3: 12}
        size = sizes.get(level, 12)
        for tok, b, mono in strip_md_inline(text):
            if mono:
                run = h.add_run(tok)
                run.font.size = Pt(max(size - 1, 10))
                run.font.color.rgb = C_BLUE
                run.font.name = 'Consolas'
                set_east(run, '等线')
            else:
                run = h.add_run(tok)
                run.font.size = Pt(size)
                run.font.color.rgb = C_BLUE
                run.bold = True
                run.font.name = 'Calibri'
                set_east(run, '黑体')
        h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        h.paragraph_format.space_after = Pt(6)
        return h

    def add_table_block(headers, rows):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'
        for j, htext in enumerate(headers):
            cell = t.cell(0, j)
            cell.paragraphs[0].text = ''
            p = cell.paragraphs[0]
            for tok, b, mono in strip_md_inline(htext):
                add_run(p, tok, east='黑体', size=10, bold=True,
                        color=RGBColor(0xFF, 0xFF, 0xFF), mono=mono)
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), '2F5B8C')
            cell._element.get_or_add_tcPr().append(shd)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = t.cell(i + 1, j)
                cell.paragraphs[0].text = ''
                p = cell.paragraphs[0]
                for tok, b, mono in strip_md_inline(val):
                    add_run(p, tok, east='宋体', size=10, bold=b or None, mono=mono)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
        return t

    def add_code_block(text):
        for ln in text.split('\n'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.2)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)
            add_run(p, ln if ln else ' ', east='等线', size=9, mono=True, color=C_CODE)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)

    def add_mermaid_image(idx):
        png = os.path.join(png_dir, f'mermaid_{idx}.png')
        if not os.path.exists(png):
            add_para(f'[图表渲染失败: {idx}]', size=10, color=C_GRAY)
            return
        img = Image.open(png)
        w, h = img.size
        aspect = h / w if w else 1
        target_w = MAX_W
        target_h = target_w * aspect
        if target_h > MAX_H:
            target_h = MAX_H
            target_w = target_h / aspect
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(png, width=Inches(target_w), height=Inches(target_h))

    # 文档标题
    if title:
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_after = Pt(2)
        add_run(title_p, title, east='黑体', size=22, bold=True, color=C_BLUE)
    if subtitle:
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_p.paragraph_format.space_after = Pt(12)
        add_run(sub_p, subtitle, east='黑体', size=18, bold=True, color=C_BLUE)

    mermaid_idx = 0
    i = 0
    in_code = False
    code_buf = []
    code_lang = ''

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_buf = []
                if code_lang == 'mermaid':
                    add_mermaid_image(mermaid_idx)
                    mermaid_idx += 1
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith('```'):
                        i += 1
                    in_code = False
                    i += 1
                    continue
                i += 1
                continue
            else:
                add_code_block('\n'.join(code_buf))
                in_code = False
                i += 1
                continue

        if in_code:
            code_buf.append(raw)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == '---':
            i += 1
            continue

        if re.match(r'^\s*[-*]\s+\[.*\]\(#', stripped):
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1:
                i += 1
                continue
            if text.strip() in ('目录', 'Ŀ¼'):
                i += 1
                while i < len(lines):
                    nxt = lines[i].strip()
                    if nxt.startswith('## ') or nxt.startswith('### '):
                        break
                    if not re.match(r'^\s*[-*]\s+\[.*\]\(#', nxt):
                        break
                    i += 1
                continue
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
            add_heading(text, min(level, 3))
            i += 1
            continue

        if stripped.startswith('|'):
            headers = None
            rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                cells = [c.strip() for c in lines[j].strip().strip('|').split('|')]
                if all(re.match(r'^:?-{2,}:?$', c.replace(' ', '')) for c in cells if c):
                    j += 1
                    continue
                if headers is None:
                    headers = cells
                else:
                    rows.append(cells)
                j += 1
            if headers:
                add_table_block(headers, rows)
                i = j
                continue
            add_para([(stripped, False, False)])
            i += 1
            continue

        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            if re.match(r'^\*\*版本\*\*|^\*\*日期\*\*|^\*\*状态\*\*|^\*\*主要变更\*\*', text):
                i += 1
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)
            for tok, b, mono in strip_md_inline(text):
                add_run(p, tok, east='宋体', size=10, bold=b or None, color=C_GRAY, mono=mono)
            i += 1
            continue

        m = re.match(r'^\s*[-*]\s+(.*)', stripped)
        if m:
            text = m.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            for tok, b, mono in strip_md_inline(text):
                add_run(p, tok, east='宋体', size=10.5, bold=b or None, mono=mono)
            i += 1
            continue

        m = re.match(r'^\s*(\d+)\.\s+(.*)', stripped)
        if m:
            num = m.group(1)
            text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_after = Pt(3)
            add_run(p, num + '. ', east='宋体', size=10.5, bold=True)
            for tok, b, mono in strip_md_inline(text):
                add_run(p, tok, east='宋体', size=10.5, bold=b or None, mono=mono)
            i += 1
            continue

        add_para(strip_md_inline(stripped), size=10.5)
        i += 1

    os.makedirs(os.path.dirname(os.path.abspath(docx_path)), exist_ok=True)
    doc.save(docx_path)
    print('DOCX saved:', docx_path)
    print('paras:', len(doc.paragraphs), 'tables:', len(doc.tables))
    print('mermaid images used:', mermaid_idx)


def main():
    ap = argparse.ArgumentParser(description='Markdown + Mermaid PNG → Word 文档')
    ap.add_argument('md_path', help='markdown 文件路径')
    ap.add_argument('--docx', default=None, help='输出 docx 路径（默认同目录同名 .docx）')
    ap.add_argument('--png-dir', default=None, help='mermaid PNG 目录（默认 ./render_out）')
    ap.add_argument('--title', default='', help='文档标题（默认取 md 首个 H1）')
    ap.add_argument('--subtitle', default='', help='文档副标题')
    args = ap.parse_args()

    md_path = os.path.abspath(args.md_path)
    if not args.docx:
        args.docx = os.path.splitext(md_path)[0] + '.docx'
    if not args.png_dir:
        # 默认找与 md 同目录的 render_out
        cand = os.path.join(os.path.dirname(md_path), 'render_out')
        args.png_dir = cand if os.path.isdir(cand) else None

    title = args.title
    if not title:
        # 从 md 首个 H1 提取标题
        with open(md_path, encoding='utf-8') as f:
            for line in f:
                m = re.match(r'^#\s+(.*)', line.strip())
                if m:
                    title = m.group(1).strip()
                    break

    build_doc(md_path, args.docx, args.png_dir or '', title=title, subtitle=args.subtitle)


if __name__ == '__main__':
    main()
